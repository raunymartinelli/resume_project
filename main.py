from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

# Create a new Word document
doc = Document()

# Define styles for the document
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Adding title - Name
title = doc.add_paragraph()
title_run = title.add_run('Layse Sales Muniz Martinelli\n')
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Adding contact information
contact_info = doc.add_paragraph()
contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_info_run = contact_info.add_run('140 Erskine Av. Toronto CA- M4P 1Z2 | Phone: +1647 606 8462 | laysemartinelli@gmail.com\n\n')
contact_info_run.italic = True
contact_info_run.font.size = Pt(11)

# Adding section title
doc.add_paragraph('SUMMARY OF QUALIFICATIONS', style='Heading 2')

# Adding content to Summary of Qualifications
summary = [
    "Professional with over 17 years of experience in the Commercial, Sales and Operational Management areas, Guest service expertise, working with market development and mapping, analyzing trends and opportunities for bringing in new products and services, upkeeping the client portfolio, prospecting new commercial partners, planning strategic marketing and positioning actions.",
    "Analyzing performance indicators and management reports, readjusting business strategies and approaches for meeting the established goals.",
    "Working alongside the senior managers and other companies' internal departments, giving qualification training to collaborators and elaborating results-oriented strategic alignments.",
    "Directly participating in the implementation of hotel, structuring and dimensioning teams, analyzing internal routines for reducing operational costs, increasing efficiency and service quality."
]

for point in summary:
    doc.add_paragraph(f' {point}', style='List Bullet')

# Adding Complementary Training and Languages section
doc.add_paragraph('\nCOMPLEMENTARY TRAINING AND LANGUAGES', style='Heading 2')
training = doc.add_paragraph()
training.add_run('Computing: ').bold = True
training.add_run('Microsoft Office and Zoom.\n')
training.add_run('Languages: ').bold = True
training.add_run('Native Portuguese, Fluent English, Fluent Spanish and Advanced French.\n\n')

# Adding Professional Experience section
doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 2')

# Adding jobs
jobs = [
    {
        'title': 'Event Specialist',
        'company': 'ShowTech Power and Lighting/ GES Canada LTDA - Toronto - CA',
        'points': [
            "Responsible for managing event timelines regarding power, lighting, and mechanical requests.",
            "Elaborating quotes for show manager and exhibitors.",
            "Coordinating exhibitor’s services (banner hanging, housekeeping, production, parking, and ICT)",
            "On-Site collection.",
            "Plot Electrical, lighting, and mechanical floor plan.",
            "Processing payments and invoices.",
            "Post-event recap and job summary."
        ]
    },

    {
        'title': 'Bilingual Customer Service Representative',
        'company': 'TTEC- Toronto - CA',
        'duration': 'Dec/2021 until Jul/2022',

        'points': [
            "Answer incoming communications from Booking.com Partners.",
            "Conduct research to provide answers for partners to resolve their issues in concern to booking.com extranet.",
            "Offer the highest level of service possible through voice, chat or email interactions."
        ]

    },

    {
        'title': 'Sales Associate',
        'company': "Carter's Inc - Toronto - CA",
        'duration': 'Oct/2021 until Nov/2022',

        'points': [
            "Building and maintaining long-term relationships with customers by providing an exceptional shopping experience which includes greeting customers, helping, directing customers to merchandise, product suggestion, providing product information and offering positive opinions.",
            "Working together with all sales associates and the management team meetings and exceeding store sales objectives and operational standards.",
            "Processing customer sale transactions and accurately following cash handling procedures.",
            "Assisting with merchandising, marketing and maintaining visual standards.",
            "Securing company assets by following all loss prevention policies and procedures.",
            "Developing retail skills by completing company training as required.",
            "Receiving and processing company product accurately and efficiently while upholding the organization of the backroom."

        ]

    },
{

        'title': 'Event Coordinator',
        'company': 'Consultre Consultoria e Treinamentos',
        'duration': 'Sep/2018 until Jul/2021',

        'points': [
            "Responsible for organizing, welcoming and providing service to students and teachers, gathering needs and elaborating strategic alignments for developing events.",
            "Presenting and negotiating commercial proposals, being in charge of selling courses, carrying out market surveys, analyzing trends and competition, identifying new business opportunities."

        ]

    },

    {

        'title': 'Event and Hotel Fundraising Executive',
        'company': 'Natal Convention Bureau',
        'duration': 'Feb/2019 until Mar/2020',

        'points': [
            "Responsible for visiting the corporate market for bringing in events, visiting the city hotel for finding new associates and reinstating former ones, managing events' demands, providing service to domestic and global markets, organizing famtours and business rounds."

        ]

    },

    {

        'title': 'Accounts Executive',
        'company': 'Hotel Wish Natal - Rede GJP',

        'points': [
            "Carrying out market studies and analyzing business opportunities, maintaining accounts, prospecting customers, bringing in new events, managing demands and follow-ups.",
            "Developing customer relationship management routines, registering clients, managing goals, negotiating fees, making reservations, launching campaigns and elaborating strategic plans in general and leading the sales team on providing customer service."

        ]

    },

    {

        'title': 'Sales Manager',
        'company': 'Atlantica Hotels International',

        'points': [
            "Mapping the market, opening and introducing new products, planning the yearly budget of the department, managing sales and bookings reports, managing teams and events.",
            "Planning, coordinating and evaluating routines, leading corporate visits for prospecting businesses and upkeeping key accounts, leading the sales team on providing customer service."

        ]

    },

    {

        'title': 'Account Executive - Corporate and Leisure Market',
        'company': 'Rede Allia Hotels - BHG - Solare Belem',

        'points': [
            "Carrying out market studies for introducing the product Gran Solare Connext Loft & Office, visiting clients, negotiating fee arrangements, promoting famtours and events and leading the sales team on providing customer service.",
            "Utilizing Pmweb as an internal system for registering clients, supplying fees, negotiating blocks, issuing reports, performance indicators and results' analysis."

        ]

    },

    {

        'title': 'General/Commercial/Hosting/F&D Manager',
        'company': 'Araca Praia Flat',

        'points': [
            "Managing administrative and commercial routines, strategic planning marketing actions, creating internal policies, analyzing business opportunities and promoting events.",
            "Planning, evaluating and giving support on internal routines, leading the reception and reservation teams on negotiations and building relationships with the customers and leading the sales team on providing customer service",
            "Elaborating and implementing the strategic sales plan, expanding the company's client base, positioning the company within the market and highlighting the service on touristic channels.",
            "In charge of recruiting and training collaborators to compose the Commercial, Food and Drinks, Reception and Governance teams, as well as other collaborators according to internal needs.",
            "Participating in fairs, workshops and business trips to keep up to date and obtain qualification, visiting tourism agencies and the corporate market nationwide."

        ]

    }

]

for job in jobs:
    job_title = doc.add_paragraph()
    job_title.add_run(f'{job["title"]} | {job["company"]}').bold = True
    # Check if 'duration' key exists and add it on a new line
    duration = job.get('duration', '')
    if duration:
        job_title.add_run(f'\n{duration}').italic = True
    for point in job['points']:
        doc.add_paragraph(f' {point}', style='List Bullet')

# Saving the document
file_path = '/Users/martinelli/Desktop/Layse_Sales_Muniz_Martinelli_Resume.docx'

doc.save(file_path)

file_path


